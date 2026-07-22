from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\zhepeng\OneDrive\CAU\学校\经济时间序列")
OUT = ROOT / "Enders_Chapter1_中文讲义_差分方程.docx"
FIG = ROOT / "enders_figure_1_1_reconstructed.png"


def set_font(run, name="宋体", size=10.5, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_cm):
        col.set(qn("w:w"), str(int(Cm(width).twips)))
        col.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(Cm(width).twips)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_font(run, "宋体", 9, color=(100, 100, 100))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    set_font(run2, "宋体", 9, color=(100, 100, 100))


def add_math(doc, expression):
    """Write a Word math-zone paragraph. The source file stores visible symbols, never LaTeX."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mrpr = OxmlElement("m:rPr")
    mfont = OxmlElement("m:mathFont")
    mfont.set(qn("m:val"), "Cambria Math")
    mrpr.append(mfont)
    mr.append(mrpr)
    mt = OxmlElement("m:t")
    mt.text = expression
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)
    return p


def add_body(doc, text, first_line=True):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0.74) if first_line else Cm(0)
    r = p.add_run(text)
    set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)
    return p


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [16.5])
    cell = table.cell(0, 0)
    shade(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, "微软雅黑", 10.5, bold=True, color=(31, 78, 120))
    r = p.add_run(text)
    set_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, "宋体", 9.5, italic=True, color=(90, 90, 90))


def make_figure():
    shocks = [0.45,0.78,0.68,1.32,-2.05,-1.08,-0.64,0.58,0.82,0.65,0.18,-0.33,-0.21,0.73,0.54,-0.10,-0.63,-0.48,0.21,0.68,0.92,0.28,1.12,-1.28,-0.84,-0.29,-0.52,0.05,0.85,1.03,0.87,1.14,-0.71,-0.66,-0.17,-0.29,-0.61,0.36,0.71,0.49,-0.18,-0.79,-0.61,-0.10,0.44,0.88,-0.42,-0.89,0.09,0.62]
    t, trend, seasonal, irregular, total = [], [], [], [], []
    state = 0.0
    for k in range(1, 81):
        tr = 1 + 0.1 * k
        se = 1.6 * math.sin(math.pi * k / 6)
        state = 0.7 * state + (shocks[k - 1] if k <= 50 else 0)
        t.append(k); trend.append(tr); seasonal.append(se); irregular.append(state); total.append(tr + se + state)
    width, height = 1500, 1050
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 25)
        small = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 21)
        title = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 29)
    except OSError:
        font = small = title = ImageFont.load_default()
    left, right = 120, 1430
    top1, bottom1, top2, bottom2 = 85, 470, 605, 980
    def xp(v): return left + (v / 80) * (right - left)
    def ytop(v): return bottom1 - (v / 12) * (bottom1 - top1)
    def ybottom(v): return bottom2 - ((v + 4) / 12) * (bottom2 - top2)
    def axes(top, bottom, values, mapper):
        for val in values:
            y = mapper(val)
            draw.line((left, y, right, y), fill="#e0e4e8", width=1)
            draw.text((left - 48, y - 12), str(val), fill="#5e6872", font=small)
        draw.line((left, top, left, bottom), fill="#64707d", width=2)
        draw.line((left, bottom, right, bottom), fill="#64707d", width=2)
    axes(top1, bottom1, [0,2,4,6,8,10,12], ytop)
    axes(top2, bottom2, [-4,-2,0,2,4,6,8], ybottom)
    for val in range(0,81,10):
        x = xp(val)
        draw.line((x, bottom1, x, bottom1+9), fill="#64707d", width=2)
        draw.line((x, bottom2, x, bottom2+9), fill="#64707d", width=2)
        draw.text((x-13, bottom2+15), str(val), fill="#5e6872", font=small)
    draw.text((left, 28), "Hypothetical time series", fill="#1f4e79", font=title)
    draw.text((left+15, top1+15), "Observed data", fill="#1f4e79", font=font)
    draw.text((xp(50)+18, top1+15), "Forecast", fill="#1f4e79", font=font)
    draw.text((left, top2-42), "Components", fill="#1f4e79", font=title)
    draw.text((right-80, bottom2+15), "Time t", fill="#5e6872", font=small)
    def plot(values, mapper, color, start=0, dash=False, width_line=4):
        coords = [(xp(t[start + j]), mapper(val)) for j, val in enumerate(values)]
        for (x1,y1),(x2,y2) in zip(coords[:-1], coords[1:]):
            if dash:
                length = max(1, int(math.hypot(x2-x1, y2-y1)))
                for z in range(0, length, 14):
                    if (z//14) % 2 == 0:
                        a, b = z/length, min(z+8, length)/length
                        draw.line((x1+(x2-x1)*a, y1+(y2-y1)*a, x1+(x2-x1)*b, y1+(y2-y1)*b), fill=color, width=width_line)
            else:
                draw.line((x1,y1,x2,y2), fill=color, width=width_line)
    plot(total[:50], ytop, "#1f4e79")
    plot(total[49:], ytop, "#1f4e79", start=49, dash=True)
    draw.line((xp(50), top1, xp(50), bottom1), fill="#777777", width=2)
    plot(trend, ybottom, "#1f4e79")
    plot(seasonal, ybottom, "#b05a2a", dash=True)
    plot(irregular, ybottom, "#555555", width_line=3)
    draw.line((left, ybottom(0), right, ybottom(0)), fill="#9099a2", width=1)
    legend_y = top2 + 8
    for x, label, color, dashed in [(left+260,"Trend","#1f4e79",False),(left+470,"Seasonal","#b05a2a",True),(left+735,"Irregular","#555555",False)]:
        if dashed:
            draw.line((x, legend_y, x+46, legend_y), fill=color, width=3)
            draw.line((x+58, legend_y, x+104, legend_y), fill=color, width=3)
        else:
            draw.line((x, legend_y, x+104, legend_y), fill=color, width=3)
        draw.text((x+116, legend_y-13), label, fill="#4b5560", font=small)
    image.save(FIG)


def configure(doc):
    sec = doc.sections[0]
    # Named override: A4 academic handout for Chinese university teaching.
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(2.1), Cm(2.0)
    sec.left_margin, sec.right_margin = Cm(2.2), Cm(2.2)
    sec.header_distance, sec.footer_distance = Cm(1.1), Cm(1.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 24, (11, 37, 69), 0, 8),
        ("Heading 1", 16, (46, 116, 181), 16, 8),
        ("Heading 2", 13, (46, 116, 181), 12, 6),
        ("Heading 3", 11.5, (31, 77, 120), 8, 4),
    ]:
        st = styles[name]
        st.font.name = "微软雅黑"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size); st.font.color.rgb = RGBColor(*color); st.font.bold = True
        st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, "微软雅黑", 25, bold=True, color=(11, 37, 69))
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(subtitle)
        set_font(r, "微软雅黑", 14, color=(68, 90, 112))


def build():
    make_figure()
    doc = Document()
    configure(doc)

    add_title(doc, "经济时间序列讲义", "第1章  差分方程与经济动态")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向应用经济学硕士生与低年级博士生")
    set_font(r, "宋体", 12, color=(80, 80, 80))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    r = p.add_run("依据 Walter Enders《Applied Econometric Time Series》第1章重构\n中文原创教学版")
    set_font(r, "宋体", 10.5, color=(100, 100, 100))
    doc.add_page_break()

    doc.add_heading("本章导读", level=1)
    add_body(doc, "本章不是把差分方程当作纯粹的数学预备知识，而是把它视为描述经济系统跨期调整的语言。一个经济变量之所以需要时间序列模型，是因为它通常既继承过去状态，也受到当期新信息和随机冲击的影响。后续的 ARMA、单位根、VAR、协整和误差修正模型，都可以在这一框架下理解。")
    doc.add_heading("学习目标", level=2)
    for item in [
        "能够用“状态、冲击、传导、均衡和预测”解释一个动态经济模型；",
        "能够区分结构式方程与简化式方程，并说明二者的用途；",
        "掌握一阶与二阶线性差分方程的基本求解思路；",
        "理解稳定性、特征根、初始条件、特解和齐次解的经济含义；",
        "能够使用滞后算子简洁地表示动态模型；",
        "能够判断后顾型与前瞻型模型在求解逻辑上的差别。",
    ]: add_bullet(doc, item)
    doc.add_heading("建议授课安排", level=2)
    table = doc.add_table(rows=1, cols=3); set_table_widths(table, [2.4, 7.2, 6.9])
    headers = ["课时", "内容", "课堂重点"]
    for c, x in zip(table.rows[0].cells, headers):
        shade(c, "E8EEF5"); r = c.paragraphs[0].add_run(x); set_font(r, "微软雅黑", 10, bold=True)
    for row in [
        ("第1次", "1.1 时间序列模型；随机游走；结构式与简化式；误差修正", "动态建模为何是时间序列分析的起点"),
        ("第2次", "1.2 差分方程及其解；迭代法；稳定性；蛛网模型", "冲击会消失、持续还是放大"),
        ("第3次", "高阶方程、特解、滞后算子、前瞻与后顾型解", "从经济模型走向可估计的时间序列形式"),
    ]:
        cells = table.add_row().cells
        for c, x in zip(cells, row):
            r = c.paragraphs[0].add_run(x); set_font(r, "宋体", 9.5)

    doc.add_heading("1.1 时间序列模型", level=1)
    add_body(doc, "时间序列是按时间顺序记录的同一变量的一组观测值，通常写作 {yₜ}。它与横截面数据的根本差异在于：观测值的先后顺序本身包含信息。今天的通胀、价格、投资和产出往往依赖过去的状态；一个意外冲击也可能通过滞后关系持续影响未来。")
    add_note(doc, "核心观点", "时间序列模型首先是一套关于“冲击如何跨期传导”的经济动态模型；统计估计与预测建立在这一动态结构之上。")
    doc.add_heading("一、分解视角：趋势、季节与不规则波动", level=2)
    add_body(doc, "传统预测方法常将观测序列理解为若干成分的叠加。最简单的加法分解可写为：")
    add_math(doc, "yₜ = Tₜ + Sₜ + Iₜ")
    add_body(doc, "其中，Tₜ 为趋势成分，刻画长期水平或长期增长；Sₜ 为季节成分，刻画固定周期的规律性波动；Iₜ 为不规则成分，刻画未被趋势和季节解释的短期扰动。现代观点进一步强调：三个部分都可能含有随机性，因此分解本身也是一个需要估计的问题。")
    doc.add_picture(str(FIG), width=Cm(16.0))
    add_caption(doc, "图1.1  假想时间序列的重建图。趋势、季节和不规则项按教材给出的运动方程生成；未来期令新扰动为零。")
    add_body(doc, "图中的具体设定为：趋势项以固定速度增长，季节项每12期循环，不规则项则具有正相关性。若观察到第50期，短期预测仍会利用当前不规则项的惯性；但随着预测期延长，其影响会逐渐消退。")
    add_math(doc, "Tₜ = 1 + 0.1t   ；    Sₜ = 1.6 sin(tπ/6)    ；    Iₜ = 0.7Iₜ₋₁ + εₜ")
    doc.add_heading("二、差分方程的含义", level=2)
    add_body(doc, "一般而言，差分方程把变量当前值写成其自身滞后值、时间以及其他变量的函数。例如，Iₜ = 0.7Iₜ₋₁ + εₜ 表示：本期不规则项由上一期不规则项的70%与本期不可预期冲击共同决定。系数0.7刻画冲击的持续性。")
    add_body(doc, "时间序列计量经济学关心的并非任意差分方程，而是包含随机成分的差分方程。研究者既可以估计单变量序列的规律，也可以研究多个相互作用的经济变量。")

    doc.add_heading("三、四类经济动态模型", level=2)
    doc.add_heading("1. 随机游走与有效市场直觉", level=3)
    add_body(doc, "若市场中不存在可利用的系统性预期收益，股票价格对数的变化可被近似为纯创新：")
    add_math(doc, "yₜ₊₁ = yₜ + εₜ₊₁    ⇔    Δyₜ₊₁ = εₜ₊₁")
    add_body(doc, "在更一般的检验式中，随机游走要求截距和水平项均为零：")
    add_math(doc, "Δyₜ₊₁ = a₀ + a₁yₜ + εₜ₊₁    ，随机游走的约束为 a₀ = a₁ = 0")
    add_body(doc, "这同时要求在第t期的信息集下，下一期创新的条件均值为零。注意：随机游走并不等于价格没有波动，而是指价格变化缺少可以稳定利用的可预测成分。")
    doc.add_heading("2. 结构式与简化式：Samuelson 乘数—加速数模型", level=3)
    add_math(doc, "yₜ = cₜ + iₜ")
    add_math(doc, "cₜ = αyₜ₋₁ + εcₜ    ，0 < α < 1")
    add_math(doc, "iₜ = β(cₜ − cₜ₋₁) + εiₜ    ，β > 0")
    add_body(doc, "第一式是收入恒等式；第二式表示消费随滞后收入而调整；第三式表示投资受消费变化驱动。由于投资方程中的当期消费 cₜ 也是内生变量，第三式属于结构式方程。")
    add_body(doc, "将消费方程及其滞后形式代入，可以把系统化为仅含GDP自身滞后项的单变量简化式：")
    add_math(doc, "yₜ = α(1 + β)yₜ₋₁ − αβyₜ₋₂ + xₜ")
    add_body(doc, "简化式的价值在于预测：只要估计出参数，并观测到过去的GDP，就能递推预测未来GDP。不过，简化式本身不自动揭示结构参数的经济含义；从简化式恢复结构机制需要额外的识别限制。")
    doc.add_heading("3. 误差修正：即期与远期汇率", level=3)
    add_math(doc, "sₜ₊₁ = fₜ + εₜ₊₁")
    add_body(doc, "无偏远期汇率假说认为，当前远期汇率 fₜ 是下一期即期汇率 sₜ₊₁ 的无偏预测。当即期与远期价格偏离长期均衡时，之后的变动会部分纠正这一偏离：")
    add_math(doc, "sₜ₊₂ = sₜ₊₁ − α(sₜ₊₁ − fₜ) + εs,ₜ₊₂    ，α > 0")
    add_math(doc, "fₜ₊₁ = fₜ + β(sₜ₊₁ − fₜ) + εf,ₜ₊₁    ，β > 0")
    add_body(doc, "括号中的差额是上一期的均衡误差。若 sₜ₊₁ 高于 fₜ，则模型预测即期价格下调、远期价格上调，从而缩小偏离。这一逻辑将在协整与误差修正模型中再次出现。")
    doc.add_heading("4. 非线性动态：正负冲击的反应可能不同", level=3)
    add_math(doc, "iₜ = β₁(cₜ − cₜ₋₁) − λₜβ₂(cₜ − cₜ₋₁) + εiₜ")
    add_body(doc, "令 λₜ 在消费下降时取1、其他时候取0，并假定 β₁ > β₂ > 0。于是，消费增加时投资对变化的反应系数为 β₁；消费减少时，反应系数为 β₁ − β₂。该设定刻画了企业可能更积极响应需求扩张、却对需求收缩调整较慢的现实。")

    doc.add_heading("1.2 差分方程及其解", level=1)
    add_body(doc, "差分方程是在离散时间中描述变量变化的方程。所谓“求解”，不是只找一个数值，而是给定初始条件和冲击序列后，找出变量在每个未来时期的完整路径。")
    doc.add_heading("一、基本形式、初始条件与稳态", level=2)
    add_math(doc, "yₜ = a + byₜ₋₁")
    add_body(doc, "这是最简单的一阶线性差分方程。参数a表示每期固定推动项，b表示上一期状态的保留比例。若系统存在稳态 y*，则在稳态中 yₜ = yₜ₋₁ = y*，故：")
    add_math(doc, "y* = a/(1 − b)    ，b ≠ 1")
    add_body(doc, "稳态不是样本均值，也不是任一期的观测值；它是模型隐含的长期位置。初始值 y₀ 决定系统从哪里出发，参数b决定它向稳态调整的方式和速度。")
    doc.add_heading("二、迭代法求解一阶方程", level=2)
    add_body(doc, "把方程逐期代入：")
    add_math(doc, "y₁ = a + by₀")
    add_math(doc, "y₂ = a + b(a + by₀) = a(1 + b) + b²y₀")
    add_math(doc, "yₜ = a(1 + b + ··· + bᵗ⁻¹) + bᵗy₀")
    add_body(doc, "利用等比数列求和，可将一般解写得更紧凑：")
    add_math(doc, "yₜ = a(1 − bᵗ)/(1 − b) + bᵗy₀ = y* + bᵗ(y₀ − y*)")
    add_note(doc, "读式方法", "第一项 y* 是长期位置；第二项 bᵗ(y₀ − y*) 是初始偏离在第t期仍残留的部分。它直接揭示了初始条件的影响何时消失。")
    doc.add_heading("三、稳定性与经济解释", level=2)
    table = doc.add_table(rows=1, cols=3); set_table_widths(table, [3.3, 4.8, 8.4])
    for c, x in zip(table.rows[0].cells, ["参数条件", "动态表现", "经济解释"]):
        shade(c, "E8EEF5"); r=c.paragraphs[0].add_run(x); set_font(r,"微软雅黑",10,bold=True)
    rows = [
        ("0 < b < 1", "单调收敛", "冲击逐期衰减，变量从一侧回到稳态。"),
        ("−1 < b < 0", "振荡收敛", "变量在稳态两侧来回调整，但幅度逐渐缩小。"),
        ("b = 1", "不回归稳态", "固定冲击会永久累积；加入随机项后是随机游走的基本形式。"),
        ("|b| > 1", "发散", "初始偏离和冲击不断被放大，线性系统不稳定。"),
    ]
    for row in rows:
        cells=table.add_row().cells
        for c,x in zip(cells,row): r=c.paragraphs[0].add_run(x); set_font(r,"宋体",9.5)
    doc.add_heading("四、加入随机扰动：从差分方程到AR(1)", level=2)
    add_math(doc, "yₜ = a + byₜ₋₁ + εₜ")
    add_body(doc, "现在 εₜ 表示当期无法由过去信息预见的冲击。不断迭代后：")
    add_math(doc, "yₜ = aΣⱼ₌₀ᵗ⁻¹bʲ + bᵗy₀ + Σⱼ₌₀ᵗ⁻¹bʲεₜ₋ⱼ")
    add_body(doc, "当 |b| < 1 且样本足够长时，初始条件的影响趋于零，可写为：")
    add_math(doc, "yₜ = a/(1 − b) + εₜ + bεₜ₋₁ + b²εₜ₋₂ + ···")
    add_body(doc, "该表达式说明：当前值是长期均值加上当前和全部过去冲击的加权和。越久远的冲击权重越小。这正是平稳AR过程、冲击响应和预测递推的共同基础。")

    doc.add_heading("1.3 迭代、预测与冲击传导", level=1)
    doc.add_heading("一、条件预测", level=2)
    add_body(doc, "若已知第t期信息，且未来创新的条件期望为零，一步预测为：")
    add_math(doc, "Eₜ(yₜ₊₁) = a + byₜ")
    add_body(doc, "h步预测为：")
    add_math(doc, "Eₜ(yₜ₊ₕ) = y* + bʰ(yₜ − y*)")
    add_body(doc, "因此，短期预测保留当前状态的信息；预测期拉长后，预测值逐渐靠近长期均值。一个模型若声称长期预测仍高度依赖今天的偶然波动，就隐含其过程存在很强的持久性，甚至可能非平稳。")
    doc.add_heading("二、脉冲响应的最简单形式", level=2)
    add_body(doc, "若第t期出现一个单位正冲击，且模型为 yₜ = a + 0.7yₜ₋₁ + εₜ，则该冲击对各期的影响为：")
    add_math(doc, "第t期：1    ；    第t+1期：0.7    ；    第t+2期：0.49    ；    第t+3期：0.343")
    add_body(doc, "VAR模型中的脉冲响应函数只是这一逻辑在多变量系统中的推广：研究一个变量的意外变化如何经由动态关系影响其他变量。")

    doc.add_heading("1.4 特征根法与高阶差分方程", level=1)
    add_body(doc, "对高阶差分方程，逐期迭代会越来越繁琐。特征根法提供一种系统化工具。考虑二阶齐次方程：")
    add_math(doc, "yₜ = a₁yₜ₋₁ + a₂yₜ₋₂")
    add_body(doc, "尝试令解具有指数形式 yₜ = rᵗ。代入得到特征方程：")
    add_math(doc, "r² − a₁r − a₂ = 0")
    add_body(doc, "若两个特征根 r₁、r₂ 不相等，则一般解为：")
    add_math(doc, "yₜ = A₁r₁ᵗ + A₂r₂ᵗ")
    add_body(doc, "常数 A₁、A₂ 由两个初始条件确定。稳定性要求所有特征根的模小于1。若特征根为负数，路径会交替波动；若为共轭复根，路径会呈现阻尼周期；若某根的模大于1，系统会发散。")
    add_note(doc, "与后续课程的联系", "AR(2)平稳性、VAR稳定性、单位根检验中的根，以及协整系统的动态特征，本质上都在研究特征根的位置。")

    doc.add_heading("1.5 蛛网模型：供给滞后如何造成价格波动", level=1)
    add_body(doc, "农业和初级产品市场常存在生产决策滞后。生产者依据上一期价格安排本期供给，而消费者的当期需求取决于当期价格。设：")
    add_math(doc, "需求：Qᵈₜ = α − βPₜ    ，β > 0")
    add_math(doc, "供给：Qˢₜ = γ + δPₜ₋₁    ，δ > 0")
    add_body(doc, "市场出清 Qᵈₜ = Qˢₜ 后可得到价格差分方程：")
    add_math(doc, "Pₜ = (α − γ)/β − (δ/β)Pₜ₋₁")
    add_body(doc, "这里的滞后系数为 −δ/β。由于它通常为负，价格会在均衡附近上下摆动。若 δ/β < 1，波动逐步减弱；若 δ/β > 1，波动会扩大。该模型提醒我们：价格振荡不一定来自外生冲击，也可能来自经济主体面对信息与生产滞后的理性调整。")

    doc.add_heading("1.6 非齐次方程、特解与趋势", level=1)
    add_body(doc, "当差分方程含有常数、时间趋势或外生变量时，完整解由“齐次解 + 特解”构成。以一阶方程 yₜ = a + byₜ₋₁ 为例，齐次部分 yₜ = byₜ₋₁ 给出 bᵗC；常数a产生的特解是稳态 a/(1−b)。")
    add_body(doc, "若模型含有确定性趋势，例如：")
    add_math(doc, "yₜ = a + ct + byₜ₋₁")
    add_body(doc, "通常可猜测一个线性特解 A + Bt，并代回原式确定A、B。这里的关键不是记住某个公式，而是理解：外生项的形式决定特解的形式；齐次部分决定初始偏离和冲击如何衰减或放大。")
    doc.add_heading("待定系数法的操作步骤", level=2)
    for item in [
        "先写出齐次差分方程，并求其通解；",
        "根据外生项的形式猜测特解的形式：常数对应常数，线性时间趋势对应线性函数，指数项对应同类指数函数；",
        "若猜测形式与齐次解重复，需要乘以t以避免重复；",
        "将候选特解代回原方程，比较系数后求出未知参数；",
        "用初始条件确定齐次解中的常数。",
    ]: add_number(doc,item)

    doc.add_heading("1.7 滞后算子：压缩动态结构的记号", level=1)
    add_body(doc, "滞后算子 L 的定义为：")
    add_math(doc, "Lyₜ = yₜ₋₁    ；    L²yₜ = yₜ₋₂    ；    Lᵏyₜ = yₜ₋ₖ")
    add_body(doc, "因此，AR(2)模型可写为：")
    add_math(doc, "yₜ − a₁yₜ₋₁ − a₂yₜ₋₂ = εₜ")
    add_math(doc, "(1 − a₁L − a₂L²)yₜ = εₜ")
    add_body(doc, "记号上的好处是，它使多期滞后结构变成一个多项式。令 A(L)=1−a₁L−a₂L²，则模型写为 A(L)yₜ=εₜ。后续 ARMA 模型的平稳性与可逆性条件，正是通过这一类滞后多项式的根来表述。")
    add_note(doc, "常见误解", "L不是一个普通数值变量。它是作用于时间序列的算子；对不同对象作用时，必须保留其“向后移动时间下标”的含义。")

    doc.add_heading("1.8 后顾型与前瞻型模型", level=1)
    add_body(doc, "后顾型模型使用已经实现的滞后变量，例如 yₜ = a + byₜ₋₁ + εₜ。给定初始值后，可以自然地向前迭代。前瞻型模型则含有预期或未来变量，例如：")
    add_math(doc, "yₜ = a + bEₜ(yₜ₊₁) + εₜ")
    add_body(doc, "前瞻型模型不能只依靠初始值求解，因为当前决策取决于对未来的预期。通常还需要施加不爆炸条件或横截条件，以排除不合理的泡沫式路径。经济学上，这意味着模型必须同时满足行为方程与长期可持续性。")
    add_body(doc, "这一差异很重要：后顾型模型强调过去状态的惯性；前瞻型模型强调预期如何把未来信息带回当前决策。宏观经济中的资产定价、汇率和新凯恩斯主义模型经常包含前瞻型成分。")

    doc.add_heading("本章总结", level=1)
    for item in [
        "差分方程是时间序列模型的动态骨架：它刻画变量如何由过去状态和新冲击共同决定。",
        "稳态描述长期位置；稳定性描述初始偏离和冲击能否消失。",
        "结构式服务于经济机制解释，简化式服务于预测和统计估计；由后者恢复前者需要识别限制。",
        "随机游走、误差修正、蛛网模型和非线性调整说明，不同经济理论都能自然写成随机差分方程。",
        "迭代法、特征根法与滞后算子，是从经济模型走向ARMA、VAR、单位根和协整分析的共同工具。",
    ]: add_bullet(doc,item)

    doc.add_heading("课堂练习与讨论", level=1)
    exercises = [
        "设 yₜ = 2 + 0.6yₜ₋₁，且 y₀ = 0。计算 y₁、y₂、y₃，并求长期稳态。",
        "若 yₜ = 2 − 0.5yₜ₋₁，说明其路径为何会围绕稳态振荡。",
        "对 yₜ = yₜ₋₁ + εₜ，解释一个单位冲击对第t、t+1和t+10期的影响。",
        "从 Samuelson 模型的三条结构式出发，说明为什么投资方程不是简化式。",
        "在蛛网模型中，需求对价格越敏感意味着β越大。为什么这有助于稳定价格波动？",
        "解释“能被过去值预测”与“由过去值因果导致”之间的区别。",
    ]
    for x in exercises: add_number(doc,x)
    doc.add_heading("练习提示", level=2)
    add_body(doc, "第1题的稳态为5，且路径单调收敛；第2题的稳态为4，负的滞后系数导致交替调整；第3题中冲击永久保留；第4题的关键是当期消费也是内生变量；第5题的稳定条件是 δ/β < 1；第6题应区分预测关系、格兰杰意义与结构因果识别。", first_line=False)

    doc.add_heading("符号表", level=1)
    table = doc.add_table(rows=1, cols=2); set_table_widths(table, [3.0, 13.5])
    for c,x in zip(table.rows[0].cells,["符号","含义"]):
        shade(c,"E8EEF5"); r=c.paragraphs[0].add_run(x); set_font(r,"微软雅黑",10,bold=True)
    for a,b in [
        ("yₜ", "第t期目标变量，例如产出、价格或资产价格对数。"),
        ("εₜ", "第t期随机创新或不可预期冲击，通常假设条件均值为零。"),
        ("Δ", "一阶差分算子，Δyₜ = yₜ − yₜ₋₁。"),
        ("L", "滞后算子，Lyₜ = yₜ₋₁。"),
        ("y*", "模型隐含的稳态或长期均衡位置。"),
        ("Eₜ(·)", "在第t期信息集下的条件期望。"),
        ("α、β", "经济行为关系中的参数；其经济解释取决于具体模型。"),
    ]:
        cells=table.add_row().cells
        for c,x in zip(cells,[a,b]): r=c.paragraphs[0].add_run(x); set_font(r,"宋体",9.5)

    doc.add_heading("参考文献", level=1)
    add_body(doc, "Enders, Walter. Applied Econometric Time Series. 4th ed. Wiley, 2015. 本讲义基于该书第1章的主题结构重新组织，以服务中文课堂教学；示例、解释与图形均为教学性重构。", first_line=False)
    doc.core_properties.title = "经济时间序列讲义：第1章 差分方程与经济动态"
    doc.core_properties.author = "课程教学讲义"
    doc.save(OUT)
    print("Lecture handout created.")


if __name__ == "__main__":
    build()
